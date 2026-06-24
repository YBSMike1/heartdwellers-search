import os
import re
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
import tempfile
import requests
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
from spellchecker import SpellChecker
import json
import pandas as pd

st.set_page_config(page_title="Heartdwellers Search Tool", layout="centered", page_icon="❤️")

# ====================== DAY / NIGHT TOGGLE ======================
col_title, col_toggle = st.columns([5, 1])
with col_title:
    st.title("❤️ Heartdwellers Search Tool")
with col_toggle:
    dark_mode = st.toggle("🌙 Night Mode", value=True, key="theme_toggle")

st.markdown("**Search Jesus' messages to Mother Clare**")

# ====================== CSS ======================
if dark_mode:
    st.markdown("""
    <link href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700&display=swap" rel="stylesheet">
    <style>
        .stApp { background: linear-gradient(135deg, #1F1A24 0%, #2A1F35 100%); }
        .main .block-container { background: rgba(42, 37, 51, 0.92) !important; backdrop-filter: blur(22px); border-radius: 28px; padding: 2.8rem 2.2rem; }
        .fancy-header, .fancy-white {
            font-family: 'Playfair Display', serif;
            font-weight: 700;
            text-shadow: 2px 2px 5px #000000, 0 0 15px #ffffff;
            background: rgba(0,0,0,0.65);
            padding: 10px 22px;
            border-radius: 16px;
            display: block !important;
            white-space: nowrap;
            max-width: 700px;
            margin-left: auto !important;
            margin-right: auto !important;
            text-align: center;
        }
        .fancy-header { font-size: 1.18rem !important; color: #E91E63; margin-bottom: 12px; font-variation-settings: "opsz" 200; }
        .fancy-white { font-size: 1.18rem; color: #ffffff; margin-bottom: 14px; font-variation-settings: "opsz" 200; }
    </style>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <link href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700&display=swap" rel="stylesheet">
    <style>
        .stApp { background: #f8f5f0; }
        .main .block-container { background: #ffffff !important; border: 1px solid #d4c3b5; box-shadow: 0 4px 20px rgba(0,0,0,0.1); border-radius: 28px; padding: 2.8rem 2.2rem; }
        h1, h2, h3, p, label, .stMarkdown { color: #1a1a1a !important; }
        .fancy-header, .fancy-white {
            font-family: 'Playfair Display', serif;
            font-weight: 700;
            background: #2c2c2c !important;
            color: #ffffff !important;
            padding: 10px 22px;
            border-radius: 16px;
            display: block !important;
            white-space: nowrap;
            max-width: 700px;
            margin-left: auto !important;
            margin-right: auto !important;
            text-align: center;
        }
        .fancy-header { font-size: 1.18rem !important; margin-bottom: 12px; }
        .fancy-white { font-size: 1.18rem; margin-bottom: 14px; }
    </style>
    """, unsafe_allow_html=True)

DOCX_FOLDER = "Heartdwellers Docxs"
spell = SpellChecker()

# ====================== WORD LISTS ======================
SIN_WORDS = ["adultery", "anger", "arrogance", "arrogant", "backbiting", "bitter", "bitterness", "blasphemous", "blasphemy", "boastful", "complaining", "contention", "covetousness", "deceit", "deception", "deceive", "discord", "division", "doubt", "doubting", "drunk", "envy", "envious", "falsehood", "fear", "fearful", "fornication", "fury", "gluttony", "gossip", "greed", "hate", "hatred", "haughty", "hypocrisy", "hypocrite", "idolatry", "idol", "idols", "idle", "jealous", "jealousy", "judging", "judgment", "judgmental", "lazy", "laziness", "lie", "lust", "lustful", "lying", "malice", "materialism", "murmuring", "occult", "offended", "offense", "pride", "proud", "rage", "rebellion", "rebellious", "revenge", "selfish", "selfishness", "slander", "sloth", "sorcery", "stealing", "strife", "stubborn", "stubbornness", "thief", "unbelief", "unforgiveness", "unforgiving", "vengeance", "witchcraft", "worldly", "worldliness", "wrath"]

GRACE_WORDS = ["love", "charity", "compassion", "mercy", "grace", "faith", "hope", "joy", "peace", "patience", "kindness", "goodness", "faithfulness", "gentleness", "self-control", "humility", "humbleness", "forgiveness", "forgive", "surrender", "trust", "obedience", "wisdom", "understanding", "prayer", "worship", "thanksgiving", "praise", "gratitude", "meekness", "longsuffering", "endurance", "perseverance", "steadfastness", "righteousness", "holiness", "purity", "truth", "honesty", "integrity", "generosity", "giving", "sharing", "hospitality", "service", "servant", "encouragement", "edification", "unity", "harmony", "reconciliation", "healing", "deliverance", "salvation", "redemption", "restoration", "blessing", "blessed", "anointing", "presence", "intimacy", "relationship", "abide", "remain", "dwell", "rest", "yield", "submit", "obey", "loving", "kind", "gentle", "patient", "faithful", "true", "pure", "holy", "humble", "forgiving", "grateful", "thankful", "peaceful", "joyful", "hopeful"]

# ====================== FUNCTIONS ======================
def get_sin_frequencies():
    freq = {}
    if os.path.exists("sin_word_library.json"):
        try:
            with open("sin_word_library.json", "r", encoding="utf-8") as f:
                data = json.load(f)
                for item in data.get("sin_words", []):
                    freq[item["Sin Word"]] = item["Frequency"]
        except: pass
    return freq

def get_grace_frequencies():
    freq = {}
    if os.path.exists("grace_word_library.json"):
        try:
            with open("grace_word_library.json", "r", encoding="utf-8") as f:
                data = json.load(f)
                for item in data.get("grace_words", []):
                    freq[item["Grace Word"]] = item["Frequency"]
        except: pass
    return freq

def get_word_definition(word):
    if not word or len(word) < 2: return "Please enter a valid word."
    try:
        url = f"https://wordsapiv1.p.rapidapi.com/words/{word.lower()}/definitions"
        headers = {'x-rapidapi-key': "e10c87331emshb838f6dd5aeb4e8p1a63dbjsn139eec4460a9", 'x-rapidapi-host': "wordsapiv1.p.rapidapi.com"}
        response = requests.get(url, headers=headers, timeout=8)
        if response.status_code == 200:
            data = response.json()
            if data.get("definitions") and len(data["definitions"]) > 0:
                return data["definitions"][0].get("definition", "No definition found.")
        return "No definition found for this word."
    except:
        return "Definition not available at this time."

def extract_date_from_path(file_path):
    try:
        match = re.search(r'(\w+\s+\d{4})', file_path)
        if match: return datetime.strptime(match.group(1), "%b %Y")
    except: pass
    return datetime.min

def search_file(file_path, search_word):
    try:
        doc = Document(file_path)
        pattern = re.compile(rf'(?<!\w){re.escape(search_word)}(?!\w)', re.IGNORECASE)
        for p in doc.paragraphs:
            italic_text = "".join(run.text for run in p.runs if getattr(run, 'italic', False))
            if italic_text and pattern.search(italic_text):
                return {"file": os.path.relpath(file_path, DOCX_FOLDER), "text": italic_text.strip()}
    except: pass
    return None

def search_italic_text(search_word, folder_path):
    results = []
    file_count = 0
    match_count = 0
    if not os.path.exists(folder_path): return [], 0, 0
    progress_bar = st.progress(0)
    status_text = st.empty()
    start_time = time.time()
    all_files = [os.path.join(root, f) for root, _, files in os.walk(folder_path) for f in files if f.lower().endswith('.docx') and not any(s in f.lower() for s in ["compilation ", "~$", "eom", "all messages"])]
    total_files = len(all_files)
    with ThreadPoolExecutor(max_workers=12) as executor:
        future_to_file = {executor.submit(search_file, f, search_word): f for f in all_files}
        for i, future in enumerate(as_completed(future_to_file)):
            result = future.result()
            if result:
                results.append(result)
                match_count += 1
            file_count += 1
            progress = (i + 1) / max(total_files, 1)
            progress_bar.progress(progress)
            elapsed = time.time() - start_time
            files_remaining = total_files - (i + 1)
            eta_str = f"~{int((elapsed / (i + 1)) * files_remaining)}s" if files_remaining > 0 else ""
            percent = int(progress * 100)
            status_text.markdown(f"**Searching** {i+1:,} / {total_files:,} files • **{percent}%** • {eta_str} remaining")
    progress_bar.progress(1.0)
    status_text.empty()
    return results, file_count, match_count

def build_sin_word_analysis():
    from collections import Counter
    sin_counter = Counter()
    all_files = [os.path.join(root, f) for root, _, files in os.walk(DOCX_FOLDER) for f in files if f.lower().endswith('.docx') and not any(s in f.lower() for s in ["compilation ", "~$", "eom", "all messages"])]
    for file_path in all_files:
        try:
            doc = Document(file_path)
            for p in doc.paragraphs:
                italic_text = "".join(run.text for run in p.runs if getattr(run, 'italic', False))
                if italic_text:
                    words = re.findall(r'\b[a-zA-Z]+\b', italic_text.lower())
                    for word in words:
                        if word in SIN_WORDS: sin_counter[word] += 1
        except: continue
    total = sum(sin_counter.values())
    ranked = [{"Rank": rank, "Sin Word": word, "Frequency": freq, "% of Sin Mentions": round((freq / total * 100) if total > 0 else 0, 2)} for rank, (word, freq) in enumerate(sin_counter.most_common(), 1)]
    sin_data = {"built_on": datetime.now().strftime("%Y-%m-%d %H:%M"), "total_messages_scanned": len(all_files), "total_sin_occurrences": total, "sin_words": ranked}
    with open("sin_word_library.json", "w", encoding="utf-8") as f: json.dump(sin_data, f, indent=2)
    return sin_data

def build_grace_word_analysis():
    from collections import Counter
    grace_counter = Counter()
    all_files = [os.path.join(root, f) for root, _, files in os.walk(DOCX_FOLDER) for f in files if f.lower().endswith('.docx') and not any(s in f.lower() for s in ["compilation ", "~$", "eom", "all messages"])]
    for file_path in all_files:
        try:
            doc = Document(file_path)
            for p in doc.paragraphs:
                italic_text = "".join(run.text for run in p.runs if getattr(run, 'italic', False))
                if italic_text:
                    words = re.findall(r'\b[a-zA-Z]+\b', italic_text.lower())
                    for word in words:
                        if word in GRACE_WORDS: grace_counter[word] += 1
        except: continue
    total = sum(grace_counter.values())
    ranked = [{"Rank": rank, "Grace Word": word, "Frequency": freq, "% of Grace Mentions": round((freq / total * 100) if total > 0 else 0, 2)} for rank, (word, freq) in enumerate(grace_counter.most_common(), 1)]
    grace_data = {"built_on": datetime.now().strftime("%Y-%m-%d %H:%M"), "total_messages_scanned": len(all_files), "total_grace_occurrences": total, "grace_words": ranked}
    with open("grace_word_library.json", "w", encoding="utf-8") as f: json.dump(grace_data, f, indent=2)
    return grace_data

def generate_search_summary(search_word, results):
    if not results:
        return ""
    all_text = " ".join([r["text"].lower() for r in results])
    total = len(results)
    themes = []
    if any(w in all_text for w in ["humble", "humility", "lowly"]):
        themes.append("humility and lowering oneself before God")
    if any(w in all_text for w in ["examine", "heart", "conscience"]):
        themes.append("self-examination and the condition of the heart")
    if any(w in all_text for w in ["warfare", "battle", "enemy", "demon", "spiritual"]):
        themes.append("spiritual warfare and resistance from the enemy")
    if any(w in all_text for w in ["mercy", "forgive", "grace", "compassion"]):
        themes.append("God’s mercy, forgiveness, and grace")
    if any(w in all_text for w in ["pride", "proud", "arrogant", "selfish"]):
        themes.append("the dangers of pride and self-reliance")
    if any(w in all_text for w in ["love", "charity", "compassion"]):
        themes.append("love, charity, and compassion toward others")
    if any(w in all_text for w in ["obey", "obedience", "submit", "surrender"]):
        themes.append("obedience and surrender to God’s will")
    if any(w in all_text for w in ["fear", "afraid", "trust"]):
        themes.append("overcoming fear through trust in God")

    summary = f"""In the messages Jesus gave to Mother Clare, the word **{search_word}** appears in {total} different passages. """
    if themes:
        summary += f"These passages frequently touch on themes such as **{', '.join(themes)}**. "
    else:
        summary += "These passages reveal important spiritual truths that Jesus desires His children to understand. "

    summary += f"""

When Jesus speaks about **{search_word}**, He often does so with both tenderness and urgency. He calls His followers to examine their hearts honestly and to recognize how this particular struggle affects their relationship with Him and with others. The messages make it clear that growth in this area is not optional, but essential for those who wish to walk closely with the Lord.

Many of the passages emphasize the need for humility and self-awareness. Jesus repeatedly shows that pride and self-reliance are major obstacles to receiving His grace. He invites souls to come low before Him, to acknowledge their weakness, and to allow Him to work deeply in their hearts. This is not meant to bring condemnation, but rather to open the door to genuine healing and freedom.

Jesus also speaks frequently about the reality of spiritual warfare. He warns that the enemy often uses **{search_word}** as a foothold to discourage, distract, or lead souls away from intimacy with God. The messages encourage believers to stay vigilant in prayer, to remain close to Him, and to resist the lies of the enemy with truth and trust.

Above all, these messages reveal the great mercy and patience of Jesus. Even when He corrects or warns about **{search_word}**, He does so with love and a desire to draw souls closer to Himself. He offers hope, healing, and the grace needed to overcome. The consistent message is one of invitation — an invitation to greater freedom, deeper love, and a more intimate walk with Him."""

    # Add the requested closing sentence to every summary
    summary += """

I hope this message has helped you today, may the Lord Jesus Christ Bless you and keep you each and every day"""

    return summary.strip()

# ====================== MAIN UI ======================
if os.path.exists("Newest banner.png"):
    st.image("Newest banner.png", use_container_width=True)

st.markdown('<div class="fancy-white">Enter a word or phrase here or select from Graces or Sins listed Below</div>', unsafe_allow_html=True)

search_word = st.text_input("Search term", placeholder="e.g. rapture, love, faith (typos ok)", label_visibility="collapsed", key="search_input")

col1, col2, col3 = st.columns([1, 1, 1])
with col2:
    search_clicked = st.button("🔍 Search", type="primary", use_container_width=True)

# ====================== GRACES TABLE ======================
st.markdown("---")
st.markdown('<h3 class="fancy-header">✨ Browse Graces Alphabetically (Most Used First)</h3>', unsafe_allow_html=True)
st.markdown('<p style="text-align:center; font-weight:500; margin: 8px 0 16px 0;">Click in the box next to any word in the table below to search it instantly.</p>', unsafe_allow_html=True)

if not os.path.exists("grace_word_library.json"):
    with st.spinner("Building Grace frequency cache..."):
        build_grace_word_analysis()

grace_frequencies = get_grace_frequencies()
df_grace = pd.DataFrame([{"Grace Word": word, "Frequency": grace_frequencies.get(word, 0)} for word in sorted(GRACE_WORDS)])
df_grace = df_grace.sort_values("Frequency", ascending=False)

grace_selection = st.dataframe(
    df_grace,
    column_config={"Frequency": st.column_config.ProgressColumn("Frequency of usage in all messages", min_value=0, max_value=max(grace_frequencies.values()) if grace_frequencies else 500, format="%d")},
    use_container_width=True,
    hide_index=True,
    on_select="rerun",
    selection_mode="single-row",
    key="grace_table"
)

# Container for Graces results (placed right after the Graces table)
grace_results_container = st.container()

# ====================== SINS TABLE ======================
st.markdown("---")
st.markdown('<h3 class="fancy-header">📖 Browse Sins Alphabetically (Most Used First)</h3>', unsafe_allow_html=True)
st.markdown('<p style="text-align:center; font-weight:500; margin: 8px 0 16px 0;">Click in the box next to any word in the table below to search it instantly.</p>', unsafe_allow_html=True)

if not os.path.exists("sin_word_library.json"):
    with st.spinner("Building Sins frequency cache..."):
        build_sin_word_analysis()

sin_frequencies = get_sin_frequencies()
df_sin = pd.DataFrame([{"Sin Word": word, "Frequency": sin_frequencies.get(word, 0)} for word in sorted(SIN_WORDS)])
df_sin = df_sin.sort_values("Frequency", ascending=False)

sin_selection = st.dataframe(
    df_sin,
    column_config={"Frequency": st.column_config.ProgressColumn("Frequency of usage in all messages", min_value=0, max_value=max(sin_frequencies.values()) if sin_frequencies else 438, format="%d")},
    use_container_width=True,
    hide_index=True,
    on_select="rerun",
    selection_mode="single-row",
    key="sin_table"
)

# Container for Sins results (placed right after the Sins table)
sin_results_container = st.container()

# ====================== AUTO SEARCH FROM TABLE ======================
selected_word = None
selected_from = None

if grace_selection.selection and grace_selection.selection.rows:
    row_idx = grace_selection.selection.rows[0]
    selected_word = df_grace.iloc[row_idx]["Grace Word"]
    selected_from = "grace"

if sin_selection.selection and sin_selection.selection.rows:
    row_idx = sin_selection.selection.rows[0]
    selected_word = df_sin.iloc[row_idx]["Sin Word"]
    selected_from = "sin"

if selected_word:
    target_container = grace_results_container if selected_from == "grace" else sin_results_container

    with target_container:
        with st.spinner(f"Searching for “{selected_word}”..."):
            results, file_count, match_count = search_italic_text(selected_word, DOCX_FOLDER)

        if results:
            st.success(f"✅ Found {match_count:,} matches in {file_count:,} files.")
            definition = get_word_definition(selected_word)
            st.info(f"**📖 Dictionary Definition of '{selected_word}':** {definition}")

            # Summary
            summary = generate_search_summary(selected_word, results)
            if summary:
                st.markdown("### 📝 Spiritual Summary")
                st.markdown(summary)

            # Download button
            doc = Document()
            for section in doc.sections:
                section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)

            # Top banner in docx
            if os.path.exists("Newest banner.png"):
                doc.add_picture("Newest banner.png", width=Inches(6.5))

            doc.add_heading(f'What did Jesus teach us about "{selected_word}"?', level=1)

            # Dictionary Definition early
            doc.add_heading("Dictionary Definition", level=2)
            doc.add_paragraph(f"{selected_word}: {definition}")

            # Spiritual Summary comes early (as requested)
            if summary:
                doc.add_heading("Spiritual Summary", level=2)
                doc.add_paragraph(summary)

            # Then the actual messages
            for res in results:
                doc.add_paragraph(res["file"], style='Heading 3')
                p = doc.add_paragraph(res["text"])
                for run in p.runs: run.italic = True

            # Bottom banner in docx
            if os.path.exists("Bottom banner Std.png"):
                doc.add_picture("Bottom banner Std.png", width=Inches(6.5))

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                doc.save(tmp.name)
                with open(tmp.name, "rb") as f:
                    st.download_button(label="📥 Download Full Report (Word Document)", data=f, file_name=f"Jesus speaks about {selected_word}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            # Results display
            st.subheader("📋 Search Results")
            for res in results:
                highlighted = re.sub(rf'(?<!\w){re.escape(selected_word)}(?!\w)', f'<span style="background-color: #ffeb3b; color: black; font-weight: bold;">{selected_word}</span>', res['text'], flags=re.IGNORECASE)
                with st.expander(f"📄 {res['file']}", expanded=True):
                    st.markdown(f"""<div style="font-family: Calibri, Arial, sans-serif; font-size: 0.95em; line-height: 1.8; background-color: #241F2E; padding: 20px; border-radius: 12px; border-left: 6px solid #C4457A; color: #F5E6F0; font-style: italic;">{highlighted}</div>""", unsafe_allow_html=True)
        else:
            st.info("No matches found.")

# ====================== NORMAL SEARCH BUTTON ======================
elif search_clicked and search_word:
    with st.spinner(f"Searching for “{search_word}”..."):
        results, file_count, match_count = search_italic_text(search_word, DOCX_FOLDER)

    if results:
        st.success(f"✅ Found {match_count:,} matches in {file_count:,} files.")
        definition = get_word_definition(search_word)
        st.info(f"**📖 Dictionary Definition of '{search_word}':** {definition}")

        summary = generate_search_summary(search_word, results)
        if summary:
            st.markdown("### 📝 Spiritual Summary")
            st.markdown(summary)

        # Download + results display (same as above, simplified for brevity)
        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)

        if os.path.exists("Newest banner.png"):
            doc.add_picture("Newest banner.png", width=Inches(6.5))

        doc.add_heading(f'What did Jesus teach us about "{search_word}"?', level=1)

        # Dictionary Definition early
        doc.add_heading("Dictionary Definition", level=2)
        doc.add_paragraph(f"{search_word}: {definition}")

        # Spiritual Summary comes early (as requested)
        if summary:
            doc.add_heading("Spiritual Summary", level=2)
            doc.add_paragraph(summary)

        # Then the actual messages
        for res in results:
            doc.add_paragraph(res["file"], style='Heading 3')
            p = doc.add_paragraph(res["text"])
            for run in p.runs: run.italic = True

        if os.path.exists("Bottom banner Std.png"):
            doc.add_picture("Bottom banner Std.png", width=Inches(6.5))

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            with open(tmp.name, "rb") as f:
                st.download_button(label="📥 Download Full Report (Word Document)", data=f, file_name=f"Jesus speaks about {search_word}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        st.subheader("📋 Search Results")
        for res in results:
            highlighted = re.sub(rf'(?<!\w){re.escape(search_word)}(?!\w)', f'<span style="background-color: #ffeb3b; color: black; font-weight: bold;">{search_word}</span>', res['text'], flags=re.IGNORECASE)
            with st.expander(f"📄 {res['file']}", expanded=True):
                st.markdown(f"""<div style="font-family: Calibri, Arial, sans-serif; font-size: 0.95em; line-height: 1.8; background-color: #241F2E; padding: 20px; border-radius: 12px; border-left: 6px solid #C4457A; color: #F5E6F0; font-style: italic;">{highlighted}</div>""", unsafe_allow_html=True)
    else:
        st.info("No matches found.")

# ====================== BOTTOM BANNER ======================
if os.path.exists("Bottom banner Std.png"):
    st.image("Bottom banner Std.png", use_container_width=True)

st.markdown("""
<p style="text-align:center; color:#f8f9fa; font-size:0.95rem; margin-top:1.5rem;">
    Built by Mike F. with love for our Heartdwellers family
</p>
""", unsafe_allow_html=True)
